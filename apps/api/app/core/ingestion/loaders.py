"""Document loaders for all supported file formats and URL sources.

Each loader accepts a source (file path, URL string, or raw bytes) and returns
a list of LangChain Documents with pre-populated metadata. The LoaderFactory
selects the correct loader based on file extension or source type.
"""

from abc import ABC, abstractmethod
import csv
import io
import json
from pathlib import Path
from typing import Any

from langchain_core.documents import Document
import structlog

logger = structlog.get_logger(__name__)


# ── Abstract base ──────────────────────────────────────────────────────────────


class DocumentLoader(ABC):
    """Abstract base for all document loaders."""

    @abstractmethod
    def load(self, source: "str | Path | bytes", **kwargs: Any) -> list[Document]:
        """Load documents from source and return a list of LangChain Documents."""
        ...


# ── PDF ───────────────────────────────────────────────────────────────────────


class PDFLoader(DocumentLoader):
    """Loads PDF files using pypdf, one Document per page."""

    def load(self, source: "str | Path | bytes", **kwargs: Any) -> list[Document]:
        from pypdf import PdfReader

        if isinstance(source, bytes):
            reader = PdfReader(io.BytesIO(source))
            source_path = kwargs.get("filename", "uploaded.pdf")
        else:
            reader = PdfReader(str(source))
            source_path = str(source)

        total = len(reader.pages)
        docs: list[Document] = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            docs.append(
                Document(
                    page_content=text,
                    metadata={
                        "source": source_path,
                        "page_number": i + 1,
                        "total_pages": total,
                        "file_type": "pdf",
                    },
                )
            )

        logger.info("pdf_loaded", source=str(source_path), pages=len(docs))
        return docs


# ── DOCX ──────────────────────────────────────────────────────────────────────


class DOCXLoader(DocumentLoader):
    """Loads Word (.docx) files as a single Document containing all paragraph text."""

    def load(self, source: "str | Path | bytes", **kwargs: Any) -> list[Document]:
        import docx  # python-docx

        if isinstance(source, bytes):
            doc = docx.Document(io.BytesIO(source))
            source_path = kwargs.get("filename", "uploaded.docx")
        else:
            doc = docx.Document(str(source))
            source_path = str(source)

        full_text = "\n".join(para.text for para in doc.paragraphs if para.text.strip())
        docs = [
            Document(
                page_content=full_text,
                metadata={
                    "source": source_path,
                    "file_type": "docx",
                    "paragraph_count": len(doc.paragraphs),
                },
            )
        ]
        logger.info("docx_loaded", source=str(source_path))
        return docs


# ── Plain text & Markdown ──────────────────────────────────────────────────────


class TextLoader(DocumentLoader):
    """Loads plain text and Markdown files as a single Document."""

    def load(self, source: "str | Path | bytes", **kwargs: Any) -> list[Document]:
        if isinstance(source, bytes):
            text = source.decode("utf-8", errors="replace")
            source_path = kwargs.get("filename", "uploaded.txt")
        else:
            source_path = str(source)
            text = Path(source).read_text(encoding="utf-8", errors="replace")

        extension = Path(str(source_path)).suffix.lower().lstrip(".")
        docs = [
            Document(
                page_content=text,
                metadata={
                    "source": source_path,
                    "file_type": extension or "txt",
                },
            )
        ]
        logger.info("text_loaded", source=str(source_path), chars=len(text))
        return docs


# ── HTML ──────────────────────────────────────────────────────────────────────


class HTMLLoader(DocumentLoader):
    """Loads HTML files, extracting visible text via BeautifulSoup."""

    def load(self, source: "str | Path | bytes", **kwargs: Any) -> list[Document]:
        from bs4 import BeautifulSoup

        if isinstance(source, bytes):
            raw = source.decode("utf-8", errors="replace")
            source_path = kwargs.get("filename", "uploaded.html")
        else:
            source_path = str(source)
            raw = Path(source).read_text(encoding="utf-8", errors="replace")

        soup = BeautifulSoup(raw, "html.parser")
        for tag in soup(["script", "style", "head", "meta"]):
            tag.decompose()

        text = soup.get_text(separator="\n")
        title = soup.title.string if soup.title else None

        docs = [
            Document(
                page_content=text,
                metadata={
                    "source": source_path,
                    "file_type": "html",
                    "title": title,
                },
            )
        ]
        logger.info("html_loaded", source=str(source_path))
        return docs


# ── CSV ───────────────────────────────────────────────────────────────────────


class CSVLoader(DocumentLoader):
    """Loads CSV files, one Document per row formatted as key: value lines."""

    def load(self, source: "str | Path | bytes", **kwargs: Any) -> list[Document]:
        if isinstance(source, bytes):
            raw = source.decode("utf-8", errors="replace")
            source_path = kwargs.get("filename", "uploaded.csv")
        else:
            source_path = str(source)
            raw = Path(source).read_text(encoding="utf-8", errors="replace")

        reader = csv.DictReader(io.StringIO(raw))
        rows = list(reader)
        fieldnames = list(reader.fieldnames or [])

        docs: list[Document] = []
        for i, row in enumerate(rows):
            row_text = "\n".join(f"{k}: {v}" for k, v in row.items())
            docs.append(
                Document(
                    page_content=row_text,
                    metadata={
                        "source": source_path,
                        "file_type": "csv",
                        "row_index": i,
                        "columns": fieldnames,
                    },
                )
            )

        logger.info("csv_loaded", source=str(source_path), rows=len(docs))
        return docs


# ── JSON ──────────────────────────────────────────────────────────────────────


class JSONLoader(DocumentLoader):
    """Loads JSON files. List-of-dicts becomes one Document per item; objects become one doc."""

    def load(self, source: "str | Path | bytes", **kwargs: Any) -> list[Document]:
        if isinstance(source, bytes):
            raw = source.decode("utf-8", errors="replace")
            source_path = kwargs.get("filename", "uploaded.json")
        else:
            source_path = str(source)
            raw = Path(source).read_text(encoding="utf-8", errors="replace")

        data = json.loads(raw)

        if isinstance(data, list):
            docs = [
                Document(
                    page_content=json.dumps(item, ensure_ascii=False),
                    metadata={
                        "source": source_path,
                        "file_type": "json",
                        "item_index": i,
                    },
                )
                for i, item in enumerate(data)
            ]
        else:
            docs = [
                Document(
                    page_content=json.dumps(data, ensure_ascii=False, indent=2),
                    metadata={"source": source_path, "file_type": "json"},
                )
            ]

        logger.info("json_loaded", source=str(source_path), docs=len(docs))
        return docs


# ── URL ───────────────────────────────────────────────────────────────────────


class URLLoader(DocumentLoader):
    """Fetches a URL and extracts the main body text using trafilatura."""

    def load(self, source: "str | Path | bytes", **kwargs: Any) -> list[Document]:
        import trafilatura  # pyright: ignore[reportMissingImports]

        url = str(source)
        downloaded = trafilatura.fetch_url(url)
        if not downloaded:
            logger.warning("url_fetch_failed", url=url)
            return []

        text = trafilatura.extract(downloaded, include_comments=False, include_tables=True)
        if not text:
            logger.warning("url_extract_empty", url=url)
            return []

        docs = [
            Document(
                page_content=text,
                metadata={"source": url, "file_type": "url"},
            )
        ]
        logger.info("url_loaded", url=url, chars=len(text))
        return docs


# ── Factory ───────────────────────────────────────────────────────────────────

_EXTENSION_MAP: dict[str, type[DocumentLoader]] = {
    "pdf": PDFLoader,
    "docx": DOCXLoader,
    "doc": DOCXLoader,
    "txt": TextLoader,
    "md": TextLoader,
    "markdown": TextLoader,
    "rst": TextLoader,
    "html": HTMLLoader,
    "htm": HTMLLoader,
    "csv": CSVLoader,
    "json": JSONLoader,
}


class LoaderFactory:
    """Selects the correct DocumentLoader from a file extension or source type."""

    @staticmethod
    def from_extension(extension: str) -> DocumentLoader:
        """Return a loader for the given file extension (without leading dot)."""
        ext = extension.lower().lstrip(".")
        loader_cls = _EXTENSION_MAP.get(ext)
        if loader_cls is None:
            raise ValueError(f"Unsupported file extension: {ext!r}")
        return loader_cls()

    @staticmethod
    def from_path(path: "str | Path") -> DocumentLoader:
        """Return a loader inferred from the file path's extension."""
        ext = Path(str(path)).suffix.lower().lstrip(".")
        return LoaderFactory.from_extension(ext)

    @staticmethod
    def for_url() -> DocumentLoader:
        """Return the URL loader."""
        return URLLoader()

    @staticmethod
    def supported_extensions() -> list[str]:
        """Return all file extensions supported by the factory."""
        return sorted(_EXTENSION_MAP.keys())
