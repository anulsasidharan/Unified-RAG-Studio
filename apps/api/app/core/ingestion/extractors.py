"""Metadata extraction for ingested documents.

Each extractor inspects raw file content or parsed document structure to
produce a metadata dict. These dicts are merged into Document.metadata
by IngestionService after loading and preprocessing.
"""

from pathlib import Path
import re
from typing import Any

import structlog

logger = structlog.get_logger(__name__)


# ── Helpers ───────────────────────────────────────────────────────────────────


def _safe_str(value: Any) -> str | None:
    """Convert a value to a non-empty stripped string, or None."""
    if value is None:
        return None
    result = str(value).strip()
    return result if result else None


# ── Format-specific extractors ────────────────────────────────────────────────


def extract_pdf_metadata(source: "str | bytes") -> dict[str, Any]:
    """Extract PDF document information: title, author, creation date, page count."""
    import io

    from pypdf import PdfReader

    try:
        reader = (
            PdfReader(io.BytesIO(source))
            if isinstance(source, bytes)
            else PdfReader(str(source))
        )
        info = reader.metadata or {}
        return {
            k: v
            for k, v in {
                "title": _safe_str(info.get("/Title")),
                "author": _safe_str(info.get("/Author")),
                "subject": _safe_str(info.get("/Subject")),
                "created_at": _safe_str(info.get("/CreationDate")),
                "page_count": len(reader.pages),
            }.items()
            if v is not None
        }
    except Exception as exc:
        logger.warning("pdf_metadata_extraction_failed", error=str(exc))
        return {}


def extract_docx_metadata(source: "str | bytes") -> dict[str, Any]:
    """Extract DOCX core properties: title, author, description, dates."""
    import io

    import docx  # python-docx

    try:
        doc = (
            docx.Document(io.BytesIO(source))
            if isinstance(source, bytes)
            else docx.Document(str(source))
        )
        props = doc.core_properties
        # OOXML may expose dc:description; stubs only list a subset — use getattr.
        summary = getattr(props, "description", None)
        if summary is None:
            summary = getattr(props, "comments", None)
        return {
            k: v
            for k, v in {
                "title": _safe_str(props.title),
                "author": _safe_str(props.author),
                "description": _safe_str(summary),
                "created_at": str(props.created) if props.created else None,
                "modified_at": str(props.modified) if props.modified else None,
                "paragraph_count": len(doc.paragraphs),
            }.items()
            if v is not None
        }
    except Exception as exc:
        logger.warning("docx_metadata_extraction_failed", error=str(exc))
        return {}


def extract_html_metadata(html_content: str) -> dict[str, Any]:
    """Extract metadata from HTML <head> tags: title, og:*, description, author."""
    from bs4 import BeautifulSoup

    try:
        soup = BeautifulSoup(html_content, "html.parser")
        meta: dict[str, Any] = {}

        if soup.title:
            meta["title"] = soup.title.string

        for tag in soup.find_all("meta"):
            name = tag.get("name", "").lower()
            prop = tag.get("property", "").lower()
            content = tag.get("content", "")

            if name in ("description", "author", "keywords"):
                meta[name] = content
            elif prop == "og:title":
                meta.setdefault("title", content)
            elif prop == "og:description":
                meta.setdefault("description", content)
            elif name == "article:published_time":
                meta["published_at"] = content

        return {k: _safe_str(v) for k, v in meta.items() if v}
    except Exception as exc:
        logger.warning("html_metadata_extraction_failed", error=str(exc))
        return {}


def extract_section_headers(text: str) -> list[str]:
    """Return a list of section headers detected in plain text.

    Heuristic: Markdown ATX headers (# Heading) or lines in title-case/ALL-CAPS
    that are short, do not end in sentence punctuation, and are followed by a
    blank line.
    """
    headers: list[str] = []
    lines = text.splitlines()

    for i, line in enumerate(lines):
        stripped = line.strip()
        if not stripped or len(stripped) > 120:
            continue

        # Markdown ATX headers
        if re.match(r"^#{1,6}\s+.+", stripped):
            headers.append(re.sub(r"^#{1,6}\s+", "", stripped))
            continue

        # Title-case or ALL-CAPS lines without sentence-ending punctuation
        is_short = len(stripped.split()) <= 10
        ends_mid_sentence = stripped[-1] in ".?!,;"
        is_prominent = stripped.istitle() or stripped.isupper()
        followed_by_blank = i + 1 >= len(lines) or not lines[i + 1].strip()

        if (
            stripped[0].isupper()
            and is_short
            and not ends_mid_sentence
            and is_prominent
            and followed_by_blank
        ):
            headers.append(stripped)

    return headers


def extract_url_metadata(url: str, html_content: str | None = None) -> dict[str, Any]:
    """Extract metadata for a URL-sourced document."""
    meta: dict[str, Any] = {"source_url": url}
    if html_content:
        meta.update(extract_html_metadata(html_content))
    return meta


def extract_file_metadata(path: "str | Path") -> dict[str, Any]:
    """Extract filesystem-level metadata: filename, extension, file size."""
    p = Path(str(path))
    meta: dict[str, Any] = {
        "filename": p.name,
        "file_extension": p.suffix.lower().lstrip("."),
    }
    if p.exists():
        meta["file_size_bytes"] = p.stat().st_size
    return meta
